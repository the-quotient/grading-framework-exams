""" Generates reports for the oral exams in church history. """
""" I have a csv list called "list.csv" containing the following columns: surname, first name, id number, study profile, focus topic of the exam, grade. I want to create reports prefilled with this data. The reports are docx files with file names following this pattern: "map6_surname,first_name,id_number.docx". There is a template for these reports called "template.docx". """

import pandas as pd
import re
from docx import Document
import locale

def generate_reports(csv_path, template_path, output_folder):
    # Read CSV data into a Pandas DataFrame
    df = pd.read_csv(csv_path)

    # Load the template document
    template_doc = Document(template_path)

    # Convert the date
    df['original_date'] = df['date']
    
    locale.setlocale(locale.LC_TIME, 'de_DE.utf-8')
    # Do the conversion:
    df['date'] = pd.to_datetime(df['date'], format='%Y-%m-%d', errors='coerce')
    df['date'] = df['date'].dt.strftime('%A, %d. %B %Y')
    # Reset the locale to the default
    locale.setlocale(locale.LC_TIME, '')

    # Iterate over each row in the DataFrame
    for index, row in df.iterrows():
        # Create a new document based on the template
        report_doc = Document(template_path)

        # Replace placeholders in the document with actual data
        replace_placeholders(report_doc, row)

        # Save the report with a specific filename
        output_filename = f"{output_folder}/MAP6-{row['surname']},_{row['first_name']},_{row['id_number']},_{row['original_date']},_Protokoll.docx"
        report_doc.save(output_filename)

def generate_schedule(csv_path, output_folder):
    df = pd.read_csv(csv_path)

    output = df[['date', 'begin', 'end', 'surname', 'first_name', 'id_number', 'study_profile', 'focus_topic', 'second_topic', 'third_topic']]

    output.to_excel(f"{output_folder}/Prüfungsliste.xlsx")
    
def replace_placeholders(doc, data_row):
    # Replace placeholders in the document with actual data
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in data_row.items():
                placeholder = f"{{{key}}}"
                run.text = run.text.replace(placeholder, str(value))
    # Eliminate remaining "nan" values:
    for paragraph in doc.paragraphs:
        if ": nan" in paragraph.text:            
            paragraph.text = re.sub(': nan', ": ", paragraph.text)

def main():
    csv_path = 'data/list.csv'
    template_path = 'data/template.docx'
    output_folder = 'reports'

    # Create the output folder if it doesn't exist
    import os
    os.makedirs(output_folder, exist_ok=True)

    generate_reports(csv_path, template_path, output_folder)
    generate_schedule(csv_path, output_folder)

if __name__ == "__main__":
    main()
